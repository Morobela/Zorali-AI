"""File format expansion (Phase 5): .docx and .xlsx ingestion.

Fixtures are generated in-test with python-docx / openpyxl (the same
libraries the extractor uses), uploaded through the real endpoint, polled to
``ready``, and then found through /api/files/search. Also covers the
configurable MAX_UPLOAD_MB ceiling and that the extension allowlist and
path-traversal rejections are unchanged.
"""
from __future__ import annotations

import io
from uuid import uuid4

from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app

client = TestClient(app)


def _project(name: str) -> dict:
    return client.post("/api/project", json={"name": f"{name}-{uuid4().hex[:6]}"}).json()


def _docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("The Andromeda proposal ships in October.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Milestone"
    table.rows[0].cells[1].text = "Owner"
    table.rows[1].cells[0].text = "Prototype"
    table.rows[1].cells[1].text = "Dana"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws.append(["Item", "Cost"])
    ws.append(["Zeppelin rental", 1200])
    ws.append(["Catering", 300])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _upload_and_wait(project_id: str, filename: str, payload: bytes, content_type: str) -> dict:
    res = client.post(
        f"/api/files/upload?project_id={project_id}",
        files={"file": (filename, payload, content_type)},
    )
    assert res.status_code == 202, res.text
    body = res.json()
    # TestClient runs the background ingestion before returning; poll anyway
    # to exercise the documented status flow.
    status = client.get(f"/api/files/{body['id']}/status").json()
    assert status["indexing_status"] == "ready"
    return body


def test_docx_upload_indexes_and_is_searchable():
    p = _project("fmt-docx")
    _upload_and_wait(
        p["id"], "proposal.docx", _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    hits = client.get(f"/api/files/search?project_id={p['id']}&q=Andromeda proposal").json()
    assert hits, "docx content should be retrievable"
    assert "Andromeda" in hits[0]["text"]
    # Table rows survive as tab-separated lines.
    listed = client.get(f"/api/files/list?project_id={p['id']}").json()
    all_text = " ".join(c["text"] for c in listed[0]["chunks"])
    assert "Prototype\tDana" in all_text


def test_xlsx_upload_indexes_with_sheet_headers_and_rows():
    p = _project("fmt-xlsx")
    _upload_and_wait(
        p["id"], "budget.xlsx", _xlsx_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    hits = client.get(f"/api/files/search?project_id={p['id']}&q=Zeppelin rental").json()
    assert hits, "xlsx content should be retrievable"
    listed = client.get(f"/api/files/list?project_id={p['id']}").json()
    all_text = " ".join(c["text"] for c in listed[0]["chunks"])
    assert "# Sheet: Budget" in all_text
    assert "Zeppelin rental\t1200" in all_text


def test_corrupt_docx_marks_ready_with_note_not_failed():
    p = _project("fmt-corrupt")
    body = _upload_and_wait(p["id"], "broken.docx", b"not a real docx", "application/octet-stream")
    listed = client.get(f"/api/files/list?project_id={p['id']}").json()
    record = next(f for f in listed if f["id"] == body["id"])
    assert "could not extract" in record["chunks"][0]["text"]


def test_upload_limit_is_configurable(monkeypatch):
    p = _project("fmt-limit")
    monkeypatch.setattr(settings, "max_upload_mb", 1)
    res = client.post(
        f"/api/files/upload?project_id={p['id']}",
        files={"file": ("big.txt", b"x" * (1024 * 1024 + 1), "text/plain")},
    )
    assert res.status_code == 400
    assert "max 1 MB" in res.json()["detail"]
    # Default ceiling is 25 MB — a 6 MB file (over the old 5 MB cap) passes.
    monkeypatch.setattr(settings, "max_upload_mb", 25)
    res = client.post(
        f"/api/files/upload?project_id={p['id']}",
        files={"file": ("mid.txt", b"y" * (6 * 1024 * 1024), "text/plain")},
    )
    assert res.status_code == 202


def test_allowlist_and_traversal_rejections_unchanged():
    p = _project("fmt-guard")
    res = client.post(
        f"/api/files/upload?project_id={p['id']}",
        files={"file": ("evil.exe", b"MZ", "application/octet-stream")},
    )
    assert res.status_code == 400
    assert "Unsupported file type" in res.json()["detail"]

    res = client.post(
        f"/api/files/upload?project_id={p['id']}",
        files={"file": ("../../etc/passwd.txt", b"root", "text/plain")},
    )
    assert res.status_code == 400
    assert res.json()["detail"] == "Invalid filename"
